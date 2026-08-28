import io
from unittest.mock import AsyncMock, Mock

import pytest
from docx import Document
from app.config import Settings
from app.routers.users import MAX_RESUME_UPLOAD_BYTES


@pytest.mark.asyncio
async def test_health_check_reports_reachable_database(async_client):
    response = await async_client.get("/api/health")

    assert response.status_code == 200
    assert response.json()["status"] == "healthy"
    assert response.json()["database"] == "reachable"


@pytest.mark.asyncio
async def test_deliberation_hides_internal_errors(async_client, monkeypatch):
    registration = await async_client.post(
        "/api/auth/register",
        json={
            "email": "deliberation-error@decisionos.ai",
            "password": "Password123!",
            "full_name": "Error Test",
        },
    )
    headers = {"Authorization": f"Bearer {registration.json()['access_token']}"}
    decision = await async_client.post(
        "/api/decisions",
        json={"title": "Error test", "description": "Verify safe error responses."},
        headers=headers,
    )

    monkeypatch.setattr(
        "app.routers.decisions.run_boardroom_deliberation",
        AsyncMock(side_effect=RuntimeError("database password leaked")),
    )
    response = await async_client.post(
        f"/api/decisions/{decision.json()['id']}/deliberate",
        headers=headers,
    )

    assert response.status_code == 500
    assert response.json()["detail"] == "Deliberation failed. Please try again."
    assert "database password" not in response.text


@pytest.mark.asyncio
async def test_deliberation_passes_language_to_graph(async_client, monkeypatch):
    registration = await async_client.post(
        "/api/auth/register",
        json={
            "email": "language-test@decisionos.ai",
            "password": "Password123!",
            "full_name": "Language Test",
        },
    )
    headers = {"Authorization": f"Bearer {registration.json()['access_token']}"}
    decision = await async_client.post(
        "/api/decisions",
        json={"title": "Language test", "description": "Verify localized deliberations."},
        headers=headers,
    )

    mocked_deliberation = AsyncMock(side_effect=ValueError("deliberation stopped for test"))
    monkeypatch.setattr("app.routers.decisions.run_boardroom_deliberation", mocked_deliberation)

    response = await async_client.post(
        f"/api/decisions/{decision.json()['id']}/deliberate?language=es",
        headers=headers,
    )

    assert response.status_code == 404
    assert mocked_deliberation.await_args.kwargs["language"] == "es"


@pytest.mark.asyncio
async def test_docx_resume_upload_extracts_paragraph_text(async_client, monkeypatch):
    registration = await async_client.post(
        "/api/auth/register",
        json={
            "email": "docx-test@decisionos.ai",
            "password": "Password123!",
            "full_name": "Document Test",
        },
    )
    headers = {"Authorization": f"Bearer {registration.json()['access_token']}"}

    document = Document()
    document.add_paragraph("Ada Lovelace")
    document.add_paragraph("Principal Software Engineer with distributed systems experience.")
    content = io.BytesIO()
    document.save(content)

    parser = AsyncMock(return_value={
        "current_role": "Principal Software Engineer",
        "career_goals": "Lead distributed systems strategy",
        "financial_runway_months": "12-18 months",
        "default_risk_tolerance": "moderate",
        "core_values": ["Reliability"],
        "personal_context": "Experienced systems engineer.",
        "extracted_skills": ["Python"],
        "summary": "Principal engineer profile.",
    })
    monkeypatch.setattr("app.routers.users.extract_profile_with_llm", parser)

    response = await async_client.post(
        "/api/users/upload-resume",
        headers=headers,
        files={
            "file": (
                "resume.docx",
                content.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert "Ada Lovelace" in parser.await_args.args[0]


@pytest.mark.asyncio
async def test_passwords_over_bcrypt_limit_are_rejected(async_client):
    response = await async_client.post(
        "/api/auth/register",
        json={
            "email": "long-password@decisionos.ai",
            "password": "a" * 73,
            "full_name": "Password Test",
        },
    )

    assert response.status_code == 422


@pytest.mark.asyncio
async def test_delete_decision_removes_outcome_memory(async_client, monkeypatch):
    registration = await async_client.post(
        "/api/auth/register",
        json={
            "email": "delete-memory@decisionos.ai",
            "password": "Password123!",
            "full_name": "Deletion Test",
        },
    )
    headers = {"Authorization": f"Bearer {registration.json()['access_token']}"}
    decision = await async_client.post(
        "/api/decisions",
        json={"title": "Delete memory", "description": "Delete associated vector memory."},
        headers=headers,
    )
    remove_memory = Mock()
    monkeypatch.setattr("app.routers.decisions.memory_manager.remove_decision_outcome", remove_memory)

    response = await async_client.delete(f"/api/decisions/{decision.json()['id']}", headers=headers)

    assert response.status_code == 204
    assert remove_memory.call_args.kwargs["decision_id"] == decision.json()["id"]


@pytest.mark.asyncio
async def test_copilot_requires_authentication(async_client):
    response = await async_client.post(
        "/api/boardroom/copilot",
        json={"query": "How should I evaluate this decision?"},
    )
    assert response.status_code == 401


@pytest.mark.asyncio
async def test_copilot_rejects_oversized_attachment_payload(async_client):
    registration = await async_client.post(
        "/api/auth/register",
        json={
            "email": "attachment-limit@decisionos.ai",
            "password": "Password123!",
            "full_name": "Attachment Test",
        },
    )
    headers = {"Authorization": f"Bearer {registration.json()['access_token']}"}
    response = await async_client.post(
        "/api/boardroom/copilot",
        headers=headers,
        json={
            "query": "Analyze this attachment",
            "attachments": [{
                "filename": "large.txt",
                "file_type": "text/plain",
                "size": 3_750_000,
                "data": "x" * 5_000_001,
            }],
        },
    )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_login_sets_an_httponly_session_cookie(async_client):
    await async_client.post(
        "/api/auth/register",
        json={"email": "cookie-test@decisionos.ai", "password": "Password123!", "full_name": "Cookie Test"},
    )
    response = await async_client.post(
        "/api/auth/login",
        json={"email": "cookie-test@decisionos.ai", "password": "Password123!"},
    )
    cookie = response.headers["set-cookie"].lower()
    assert "decisionos_access_token=" in cookie
    assert "httponly" in cookie


@pytest.mark.asyncio
async def test_decision_list_enforces_pagination(async_client):
    registration = await async_client.post(
        "/api/auth/register",
        json={"email": "paging-test@decisionos.ai", "password": "Password123!", "full_name": "Paging Test"},
    )
    headers = {"Authorization": f"Bearer {registration.json()['access_token']}"}
    for index in range(3):
        response = await async_client.post(
            "/api/decisions", headers=headers,
            json={"title": f"Decision {index}", "description": "A decision used to validate paging."},
        )
        assert response.status_code == 201

    response = await async_client.get("/api/decisions?offset=1&limit=1", headers=headers)
    assert response.status_code == 200
    assert len(response.json()) == 1


@pytest.mark.asyncio
async def test_resume_upload_size_is_limited(async_client):
    registration = await async_client.post(
        "/api/auth/register",
        json={
            "email": "large-upload@decisionos.ai",
            "password": "Password123!",
            "full_name": "Upload Test",
        },
    )
    headers = {"Authorization": f"Bearer {registration.json()['access_token']}"}

    response = await async_client.post(
        "/api/users/upload-resume",
        headers=headers,
        files={"file": ("large.txt", b"x" * (MAX_RESUME_UPLOAD_BYTES + 1), "text/plain")},
    )

    assert response.status_code == 413


def test_settings_do_not_use_a_predictable_secret_key():
    assert Settings().SECRET_KEY != "decisionos_super_secret_jwt_key_2026_production_grade_512"
